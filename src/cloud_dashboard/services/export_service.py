"""IGNIS Multi-Format Export Service (Phase G6).

Exports experiment results in Markdown, HTML, CSV, JSON, ZIP, and optional PDF/DOCX formats
with standardized origin metadata and graceful optional dependency degradation.
"""

import os
import csv
import json
import zipfile
import logging
from datetime import datetime, timezone
from typing import Dict, List, Optional, Any

from .repository_manager import RepositoryManager
from .report_service import ReportService

logger = logging.getLogger("export_service")

# Optional dependency imports
try:
    # pyrefly: ignore [missing-import]
    import weasyprint
    HAS_WEASYPRINT = True
except ImportError:
    HAS_WEASYPRINT = False

try:
    import docx
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


class ExportService:
    """Multi-format experiment result exporter."""

    def __init__(self, workspace_dir: Optional[str] = None, repository_manager: Optional[RepositoryManager] = None):
        self.workspace_dir = os.path.abspath(workspace_dir or os.getcwd())
        self.repo_mgr = repository_manager or RepositoryManager(workspace_dir=self.workspace_dir)
        self.report_service = ReportService(workspace_dir=self.workspace_dir)
        self.exports_cache_dir = os.path.join(self.workspace_dir, "reports", "exports")
        os.makedirs(self.exports_cache_dir, exist_ok=True)

    def get_format_capabilities(self) -> Dict[str, Any]:
        """Return format availability status matrix and optional dependency info."""
        return {
            "available": ["md", "html", "csv", "json", "zip"],
            "optional_status": {
                "pdf": {
                    "available": HAS_WEASYPRINT,
                    "reason": None if HAS_WEASYPRINT else "PDF export requires weasyprint package. Install with: pip install weasyprint"
                },
                "docx": {
                    "available": HAS_PYTHON_DOCX,
                    "reason": None if HAS_PYTHON_DOCX else "DOCX export requires python-docx package. Install with: pip install python-docx"
                }
            }
        }

    def export(self, experiment_id: str, format: str) -> str:
        """Unified public interface for exporting an experiment in a specified format.
        
        Returns the absolute file path to the generated export artifact ready for streaming.
        """
        fmt = format.lower().strip()
        detail = self.repo_mgr.get_experiment_detail(experiment_id)
        if not detail:
            raise ValueError(f"Experiment '{experiment_id}' not found in repository.")

        if fmt == "md" or fmt == "markdown":
            return self._export_markdown(experiment_id, detail)
        elif fmt == "html":
            return self._export_html(experiment_id, detail)
        elif fmt == "csv":
            return self._export_csv(experiment_id, detail)
        elif fmt == "json":
            return self._export_json(experiment_id, detail)
        elif fmt == "zip":
            return self._export_zip(experiment_id, detail)
        elif fmt == "pdf":
            return self._export_pdf(experiment_id, detail)
        elif fmt == "docx":
            return self._export_docx(experiment_id, detail)
        else:
            raise ValueError(f"Unsupported export format '{format}'. Supported formats: md, html, csv, json, zip, pdf, docx.")

    def _get_origin_metadata(self, experiment_id: str) -> Dict[str, str]:
        return {
            "system": "IGNIS Research Engine",
            "project_phase": "Phase G6 — Export & Publication",
            "experiment_id": experiment_id,
            "export_time": datetime.now(timezone.utc).isoformat(),
            "generator": "IGNIS ExportService v1.0"
        }

    def _export_markdown(self, experiment_id: str, detail: Dict[str, Any]) -> str:
        exp_dir = detail["directory_path"]
        md_file = os.path.join(exp_dir, "report.md")
        if not os.path.exists(md_file):
            md_file = os.path.join(self.workspace_dir, "results", "report.md")

        out_path = os.path.join(self.exports_cache_dir, f"report_{experiment_id}.md")
        meta = self._get_origin_metadata(experiment_id)
        header_text = f"<!-- EXPORT METADATA\nSystem: {meta['system']}\nPhase: {meta['project_phase']}\nExperiment ID: {experiment_id}\nExported At: {meta['export_time']}\n-->\n\n"

        content = ""
        if os.path.exists(md_file):
            with open(md_file, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            # Fallback generation
            self.report_service.generate_markdown(detail.get("metrics", {}), out_path)
            with open(out_path, "r", encoding="utf-8") as f:
                content = f.read()

        with open(out_path, "w", encoding="utf-8") as f:
            f.write(header_text + content)

        logger.info(f"Exported Markdown for {experiment_id} to {out_path}")
        return out_path

    def _export_html(self, experiment_id: str, detail: Dict[str, Any]) -> str:
        exp_dir = detail["directory_path"]
        html_file = os.path.join(exp_dir, "report.html")
        out_path = os.path.join(self.exports_cache_dir, f"report_{experiment_id}.html")

        if os.path.exists(html_file):
            with open(html_file, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            m_path = os.path.join(exp_dir, "metrics.json")
            r_path = os.path.join(exp_dir, "raw_results.json")
            man_path = os.path.join(exp_dir, "experiment_manifest.json")
            self.report_service.generate_html(m_path, r_path, man_path, out_path)
            with open(out_path, "r", encoding="utf-8") as f:
                content = f.read()

        meta = self._get_origin_metadata(experiment_id)
        meta_html = f'\n<!-- IGNIS EXPORT METADATA: exp_id={experiment_id} phase={meta["project_phase"]} time={meta["export_time"]} -->\n'
        
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(meta_html + content)

        logger.info(f"Exported HTML for {experiment_id} to {out_path}")
        return out_path

    def _export_csv(self, experiment_id: str, detail: Dict[str, Any]) -> str:
        out_path = os.path.join(self.exports_cache_dir, f"metrics_{experiment_id}.csv")
        meta = self._get_origin_metadata(experiment_id)

        with open(out_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            # Origin metadata rows
            writer.writerow(["# IGNIS EXPORT METADATA"])
            writer.writerow(["# System", meta["system"]])
            writer.writerow(["# Phase", meta["project_phase"]])
            writer.writerow(["# Experiment ID", experiment_id])
            writer.writerow(["# Overall Verdict", detail.get("overall_verdict", "UNKNOWN")])
            writer.writerow(["# Execution Duration (sec)", detail.get("execution_duration_sec", 0.0)])
            writer.writerow(["# Git Commit", detail.get("git_commit", "unknown")])
            writer.writerow(["# Exported At", meta["export_time"]])
            writer.writerow([])

            # Table 1: Scenarios Summary
            writer.writerow(["SCENARIO SUMMARY TABLE"])
            writer.writerow(["Scenario ID", "Verdict", "Duration (sec)", "Trial Count", "Fog Decision Latency Mean (sec)", "95% CI Low", "95% CI High"])
            for s in detail.get("scenarios", []):
                writer.writerow([
                    s.get("scenario_id"),
                    s.get("verdict"),
                    s.get("duration_sec"),
                    s.get("trial_count"),
                    s.get("latency_mean"),
                    s.get("latency_ci_low"),
                    s.get("latency_ci_high")
                ])
            writer.writerow([])

            # Table 2: Detailed Scenario Metrics
            writer.writerow(["DETAILED METRIC ROWS"])
            writer.writerow(["Scenario ID", "Metric Name", "Mean Value", "95% CI Low", "95% CI High", "Passed"])
            metrics_data = detail.get("metrics", {}).get("scenario_results", {})
            for sid, sinfo in metrics_data.items():
                for mname, mstruct in sinfo.get("metrics", {}).items():
                    val = mstruct.get("mean") if isinstance(mstruct, dict) else mstruct
                    ci = mstruct.get("confidence95", [None, None]) if isinstance(mstruct, dict) else [None, None]
                    passed = mstruct.get("passed") if isinstance(mstruct, dict) else True
                    writer.writerow([
                        sid,
                        mname,
                        val,
                        ci[0] if isinstance(ci, (list, tuple)) and len(ci) >= 1 else None,
                        ci[1] if isinstance(ci, (list, tuple)) and len(ci) >= 2 else None,
                        passed
                    ])

        logger.info(f"Exported CSV for {experiment_id} to {out_path}")
        return out_path

    def _export_json(self, experiment_id: str, detail: Dict[str, Any]) -> str:
        out_path = os.path.join(self.exports_cache_dir, f"export_{experiment_id}.json")
        meta = self._get_origin_metadata(experiment_id)

        combined_data = {
            "export_metadata": meta,
            "experiment_summary": {
                "experiment_id": detail.get("experiment_id"),
                "overall_verdict": detail.get("overall_verdict"),
                "timestamp": detail.get("timestamp"),
                "git_commit": detail.get("git_commit"),
                "trial_count": detail.get("trial_count"),
                "execution_duration_sec": detail.get("execution_duration_sec"),
                "platform": {
                    "os": detail.get("platform_os"),
                    "python": detail.get("platform_python"),
                    "docker": detail.get("platform_docker"),
                    "hostname": detail.get("hostname")
                }
            },
            "metrics": detail.get("metrics", {}),
            "manifest": detail.get("manifest", {})
        }

        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(combined_data, f, indent=2)

        logger.info(f"Exported JSON for {experiment_id} to {out_path}")
        return out_path

    def _export_zip(self, experiment_id: str, detail: Dict[str, Any]) -> str:
        out_path = os.path.join(self.exports_cache_dir, f"export_{experiment_id}.zip")
        exp_dir = detail["directory_path"]

        with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zf:
            if os.path.exists(exp_dir):
                for root, _, files in os.walk(exp_dir):
                    for file in files:
                        full_p = os.path.join(root, file)
                        rel_p = os.path.relpath(full_p, exp_dir)
                        zf.write(full_p, arcname=rel_p)
            zf.comment = f"IGNIS Experiment Export ZIP ({experiment_id})".encode("utf-8")

        logger.info(f"Exported ZIP for {experiment_id} to {out_path}")
        return out_path

    def _export_pdf(self, experiment_id: str, detail: Dict[str, Any]) -> str:
        if not HAS_WEASYPRINT:
            raise NotImplementedError("PDF export requires weasyprint package. Install with: pip install weasyprint")

        html_path = self._export_html(experiment_id, detail)
        out_path = os.path.join(self.exports_cache_dir, f"report_{experiment_id}.pdf")

        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        weasyprint.HTML(string=html_content).write_pdf(out_path)
        logger.info(f"Exported PDF for {experiment_id} to {out_path}")
        return out_path

    def _export_docx(self, experiment_id: str, detail: Dict[str, Any]) -> str:
        if not HAS_PYTHON_DOCX:
            raise NotImplementedError("DOCX export requires python-docx package. Install with: pip install python-docx")

        out_path = os.path.join(self.exports_cache_dir, f"report_{experiment_id}.docx")
        doc = docx.Document()
        meta = self._get_origin_metadata(experiment_id)

        doc.add_heading(f"IGNIS Experiment Report: {experiment_id}", level=0)
        p = doc.add_paragraph()
        p.add_run(f"System: {meta['system']}\n").bold = True
        p.add_run(f"Phase: {meta['project_phase']}\n")
        p.add_run(f"Overall Verdict: {detail.get('overall_verdict')}\n")
        p.add_run(f"Export Time: {meta['export_time']}\n")

        doc.add_heading("Scenario Summary", level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Scenario ID"
        hdr_cells[1].text = "Verdict"
        hdr_cells[2].text = "Trials"
        hdr_cells[3].text = "Mean Latency (s)"

        for s in detail.get("scenarios", []):
            row_cells = table.add_row().cells
            row_cells[0].text = str(s.get("scenario_id"))
            row_cells[1].text = str(s.get("verdict"))
            row_cells[2].text = str(s.get("trial_count"))
            row_cells[3].text = str(s.get("latency_mean", "-"))

        doc.save(out_path)
        logger.info(f"Exported DOCX for {experiment_id} to {out_path}")
        return out_path
